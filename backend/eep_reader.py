"""
Mineracao do EEP (Especificacao de Escopo do Projeto) para extrair nome do
projeto, codigo JIRA, chaves de elegibilidade e VDNs de roteamento.

Aceita PDF (formato mais comum), com fallback para .xlsx/.docx quando o EEP
foi entregue nesses formatos.
"""
import io
import re


def _minerar_texto(texto):
    nome_proj = re.search(r"Nome do Projeto:\s*(.*)", texto)
    jira_ivr = re.search(r"C[oó]digo do Projeto \(JIRA\):\s*(.*)", texto)
    if not jira_ivr:
        jira_ivr = re.search(r"\b((?:URA|IVR)-\d{5,7}(?:\s*/\s*(?:URA|IVR)-\d{5,7})?)\b", texto)

    words = re.findall(r"\b[a-zA-Z0-9_]{12,60}\b", texto)
    chaves_validas = []
    prefixos_chaves = (
        "Bases", "DDDHabilita", "BaseHabilita", "PrazoMaximo", "Segmentos",
        "Habilita", "DDDHabilitar", "Chave", "Switch",
    )
    ruido = ("projeto", "mutant", "responsavel", "formulario", "aprovacao", "documento", "versao")

    for w in words:
        if any(w.startswith(p) for p in prefixos_chaves):
            chaves_validas.append(w)
            continue
        maiusculas = sum(1 for c in w if c.isupper())
        if maiusculas >= 3 and not any(x in w.lower() for x in ruido):
            chaves_validas.append(w)

    chaves_finais = sorted(set(chaves_validas))

    vdns_minados = re.findall(r"(CHAVE\s*\d+\s*\([^)]+\))\s*-\s*[^-]+-\s*(VDN\s*\d+)", texto, re.IGNORECASE)
    if not vdns_minados:
        vdns_minados = re.findall(r"\b\d{7}\b|\b\d{2}\+base\+\d+\b", texto)

    return {
        "nome": nome_proj.group(1).strip() if nome_proj else None,
        "jira": jira_ivr.group(1).strip() if jira_ivr else None,
        "chaves": chaves_finais,
        "vdns": vdns_minados,
    }


def _minerar_pdf(pdf_bytes):
    from pdfminer.high_level import extract_text

    try:
        texto = extract_text(io.BytesIO(pdf_bytes))
    except Exception:
        texto = ""
    return _minerar_texto(texto)


def _minerar_xlsx(xlsx_bytes):
    import openpyxl

    try:
        wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    except Exception:
        return _minerar_texto("")
    partes = []
    for aba in wb.sheetnames:
        ws = wb[aba]
        for row in ws.iter_rows(values_only=True):
            for v in row:
                if v not in (None, ""):
                    partes.append(str(v))
    return _minerar_texto("\n".join(partes))


def _minerar_docx(docx_bytes):
    try:
        import docx
    except ImportError:
        return _minerar_texto("")
    try:
        doc = docx.Document(io.BytesIO(docx_bytes))
    except Exception:
        return _minerar_texto("")
    partes = [p.text for p in doc.paragraphs]
    for tabela in doc.tables:
        for row in tabela.rows:
            for cell in row.cells:
                partes.append(cell.text)
    return _minerar_texto("\n".join(partes))


def minerar_eep(eep_bytes, filename=""):
    """Detecta o formato pela extensao/assinatura e extrai os metadados do EEP."""
    nome_lower = (filename or "").lower()
    if nome_lower.endswith(".xlsx") or eep_bytes[:2] == b"PK" and nome_lower.endswith((".xlsx", ".docx")):
        if nome_lower.endswith(".docx"):
            dados = _minerar_docx(eep_bytes)
        else:
            dados = _minerar_xlsx(eep_bytes)
    elif nome_lower.endswith(".docx"):
        dados = _minerar_docx(eep_bytes)
    elif eep_bytes[:4] == b"%PDF":
        dados = _minerar_pdf(eep_bytes)
    else:
        dados = _minerar_pdf(eep_bytes)

    dados["nome"] = dados["nome"] or "Projeto sem nome identificado no EEP"
    dados["jira"] = dados["jira"] or "JIRA/IVR não identificado no EEP"
    dados["chaves"] = dados["chaves"] or []
    dados["vdns"] = dados["vdns"] or []
    return dados
